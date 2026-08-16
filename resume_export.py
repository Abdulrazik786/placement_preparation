import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

# Common resume section headers we look for to apply heading styling.
# Matched case-insensitively against a line's stripped text.
KNOWN_SECTION_HEADERS = {
    "summary", "objective", "skills", "technical skills", "projects",
    "internships", "internship", "experience", "work experience",
    "certifications", "education", "achievements", "profile",
}


def _split_into_lines(resume_text: str) -> list:
    return [line.strip() for line in resume_text.split("\n")]


def _is_header(line: str) -> bool:
    if not line:
        return False
    normalized = line.strip(":").lower()
    if normalized in KNOWN_SECTION_HEADERS:
        return True
    # Fallback: short, all-caps lines with no digits or colons are likely headers too
    # (e.g. "SKILLS", "PROJECTS") - but NOT data lines like "CGPA: 7.35" which are
    # technically "all uppercase" too since isupper() ignores digits/punctuation.
    if ":" in line or any(ch.isdigit() for ch in line):
        return False
    return line.isupper() and 2 <= len(line.split()) <= 4


def build_docx(resume_text: str) -> io.BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = _split_into_lines(resume_text)
    non_empty = [(i, l) for i, l in enumerate(lines) if l]

    # First non-empty line = name (large, bold, centered). Second, if it's not already a
    # section header, is treated as the contact/target-role line (centered, gray, smaller).
    name_idx = non_empty[0][0] if non_empty else None
    contact_idx = None
    if len(non_empty) > 1 and not _is_header(non_empty[1][1]):
        contact_idx = non_empty[1][0]

    for i, line in enumerate(lines):
        if not line:
            doc.add_paragraph("")
            continue

        if i == name_idx:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(20)
            continue

        if i == contact_idx:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            continue

        if _is_header(line):
            heading = doc.add_heading(line.upper(), level=2)
            heading.style.font.size = Pt(12.5)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(4)
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf(resume_text: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10.5, leading=14, spaceAfter=4)
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"], fontSize=12,
        spaceBefore=12, spaceAfter=2, textColor="#14213D",
    )
    name_style = ParagraphStyle(
        "Name", parent=styles["Normal"], fontSize=20, leading=24,
        alignment=TA_CENTER, textColor="#14213D", spaceAfter=2, fontName="Helvetica-Bold",
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontSize=9.5,
        alignment=TA_CENTER, textColor="#6B7280", spaceAfter=10,
    )

    lines = _split_into_lines(resume_text)
    non_empty = [(i, l) for i, l in enumerate(lines) if l]
    name_idx = non_empty[0][0] if non_empty else None
    contact_idx = None
    if len(non_empty) > 1 and not _is_header(non_empty[1][1]):
        contact_idx = non_empty[1][0]

    story = []
    for i, line in enumerate(lines):
        if not line:
            story.append(Spacer(1, 6))
            continue
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if i == name_idx:
            story.append(Paragraph(safe_line.upper(), name_style))
            continue
        if i == contact_idx:
            story.append(Paragraph(safe_line, contact_style))
            continue

        if _is_header(line):
            story.append(Paragraph(safe_line.upper(), heading_style))
            story.append(HRFlowable(width="100%", thickness=0.75, color="#D1D5DB", spaceAfter=6))
        else:
            story.append(Paragraph(safe_line, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer